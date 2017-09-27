#!/usr/bin/python
# -*- coding: utf-8 -*-

from docx import Document
import os
import re
import sys
import glob
import errno

files = glob.glob('**/*.docx')
folder = "Target"
fname = "New"
fn = "New"
reference = ""
sorted_list = []
for i in files:
    sorted_list.append(i)
    sorted_list.sort()

for name in sorted_list:
    filename = os.path.splitext(name)[0]
    print filename

    # Initialization of variables
    document = Document(name)
    table = document.tables[0]
    row_count = len(table.rows)

    # Check if directory exists
    if not os.path.exists(folder):
        os.mkdir(folder)

    print "Conversion in progress !"

    transrow = 2
    for engrow in xrange(2, row_count):

        # Read table from docx file
        row1 = table.rows[engrow]
        row2 = table.rows[transrow]


        cell = row1.cells[1]
        paragraph = cell.paragraphs[0]
        p = paragraph.text.encode('utf-8').strip()

        chunknum = re.search("====== (\d+)", p)
        if chunknum:
            if os.path.exists(folder + "/" + fname + "/" + fn):
                outfile.write("\n\n_" + reference.strip().encode("utf-8") + "_")
                print reference.encode("utf-8")
                outfile.close()
                reference = ""
            s = chunknum.group(1)
            lenstr = 0
            for i in s:
                lenstr += 1
            if lenstr < 2:
                fname = "0" + chunknum.group(1)
            else:
                fname = chunknum.group(1)
            if not os.path.exists(folder + "/" + fname):
                os.mkdir(folder + "/" + fname)
            fn = "01.md"
            outfile = open('%s/%s/%s' %(folder, fname, fn), "w")
            transrow += 1
        elif re.search("https", p):
            link = "\n\n" + p + "\n\n"
            outfile.write(link)
            # transrow += 1 # Comment this line if the target language in the source file doesnt have the links
        elif p.startswith("//"):
            if row2.cells[2].text:
                removespace = row2.cells[2].text.lstrip("\s")
                reference = row2.cells[2].text.lstrip("//")
            else:
                removespace = p.lstrip("\s")
                reference = p.lstrip("//")
            transrow += 1
        elif p.endswith("//"):
            if row2.cells[2].text:
                removespace = row2.cells[2].text.rstrip(" ")
                reference += removespace.rstrip("//")
            else:
                removespace = p.rstrip(" ")
                reference += removespace.rstrip("//")
            transrow += 1
        elif p.endswith("======"):
            if row2.cells[2].text.endswith("======="):
                heading = "# " + row2.cells[2].text.rstrip(" =======")
            elif row2.cells[2].text.endswith("======"):
                heading = "# " + row2.cells[2].text.rstrip(" ======")
            else:
                heading = "# " + row2.cells[2].text.strip()
            outfile.write(heading.encode("utf-8"))
            transrow += 1
        else:
            content = row2.cells[2].text
            outfile.write(content.encode("utf-8"))
            transrow += 1

    print "Conversion done !"

if os.path.exists(folder + "/" + fname + "/" + fn):
    outfile.write("\n\n_" + reference.strip().encode("utf-8") + "_")
    print reference.encode("utf-8")
    outfile.close()
