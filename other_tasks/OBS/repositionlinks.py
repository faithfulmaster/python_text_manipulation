#!/usr/bin/python
# -*- coding: utf-8 -*-

from docx import Document
import os
import re
import sys
import glob
import errno

files = glob.glob('**/*.docx')
for name in files:
    filename = os.path.splitext(name)[0]
    print filename

    # Initialization of variables
    document = Document(name)
    table = document.tables[0]
    row_count = len(table.rows)
    folder = "Target"

    # Check if directory exists
    if not os.path.exists(folder):
        os.mkdir(folder)

    print "Conversion in progress !"
    for r in xrange(2, row_count):

        # Read table from docx file
        row = table.rows[r]
        cell = row.cells[2]
        paragraph = cell.paragraphs[0]
        p = paragraph.text.encode('utf-8').strip()

        linksearch = re.search("https", p)
        if linksearch:
        	row.cells[1].text = p.strip()
    document.save(folder + "/" + name.split("/")[1])
    print "Conversion done !"
