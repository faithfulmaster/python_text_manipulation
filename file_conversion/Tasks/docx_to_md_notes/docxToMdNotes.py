#!/usr/bin/python
# -*- coding: utf-8 -*-

from docx import Document
import os
import re
import sys
import glob
import errno

files = glob.glob('**/*.docx')
for name in files:
    filename = os.path.splitext(name)[0]
    print filename

    # Initialization of variables
    document = Document(name)
    table = document.tables[0]
    row_count = len(table.rows)
    question = ""
    answer = ""
    chunknum = ""
    prev_chunk = ""
    prev_to_prev_chunk = ""
    pattern = re.compile(r'-')
    f = filename.split('.')
    folder = f[0]

    # Check if directory exists
    if not os.path.exists(folder):
        os.mkdir(folder)

    # Convert into md file
    print "Conversion in progress !"
    for r in xrange(1, row_count):

        # Read table from docx file
        row = table.rows[r]

        cell = row.cells[1]
        paragraph = cell.paragraphs[0]
        chapter = paragraph.text.encode('utf-8').strip()
        if len(chapter) < 2:
            chapter = '0' + chapter

        if re.search('\d+', chapter):
            if not os.path.exists('%s/%s' %(folder, chapter)):
                    os.mkdir("%s/%s/" %(folder, chapter))


        cell = row.cells[2]
        paragraph = cell.paragraphs[0]
        verse = paragraph.text.encode('utf-8').strip()
        if verse != None:
            if re.search('-',verse):
                verse = verse.split('-')[0]
                if len(verse) < 2:
                    verse = '0' + verse


        cell = row.cells[4]
        for paragraph in cell.paragraphs:
            txt = "# " + paragraph.text.encode('utf-8')
            if chapter!= None and verse != None and txt!= "":
                txt1 = txt.replace('-', '\n\n')
                txt2 = txt1.replace('–', '\n\n')
                if os.path.exists('%s/%s/%s.md' %(folder, chapter, verse)):
                    outfile = open('%s/%s/%s.md' %(folder,chapter,verse), "a")
                    outfile.write("\n" + txt2)
                    outfile.close()
                else:
                    outfile = open('%s/%s/%s.md' %(folder,chapter,verse), "w")
                    outfile.write(txt2)
                    outfile.close()

    print "Conversion done !"
