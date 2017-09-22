#!/usr/bin/python
# -*- coding: utf-8 -*-

from docx import Document
import os
import re
import sys
import glob
import errno

files = glob.glob('**/*.docx')
for name in files:
    filename = os.path.splitext(name)[0]
    print filename

    # Initialization of variables
    document = Document(name)
    table = document.tables[0]
    row_count = len(table.rows)
    question = ""
    chunknum = ""
    answer = ""
    prev_chunk = ""
    prev_to_prev_chunk = ""
    pattern = re.compile(r'-')
    f = filename.split('.')
    folder = f[0]

    # Check if directory exists
    if not os.path.exists(folder):
        os.mkdir(folder)

    # Convert into md file
    print "Conversion in progress !"
    for r in xrange(2, row_count):

        # Read table from docx file
        row = table.rows[r]
        cell = row.cells[2]
        paragraph = cell.paragraphs[0]
        p = paragraph.text.encode('utf-8').strip()

        # Check for questions
        if p.endswith('?'):
            question = p

        # Check for answers
        elif p.endswith('।') or p.endswith('"'):
            answer = p

        # Check for chunk numbers and create md files
        elif p[-1].isdigit():
            if pattern.search(p):
                chunknum = p.split('-')
                if not os.path.exists('%s/%s' %(folder, chunknum[0])):
                    os.mkdir("%s/%s/" %(folder, chunknum[0]))

                if prev_chunk == chunknum[1] or prev_to_prev_chunk == chunknum[1]:
                    outfile = open('%s/%s/%s.md' %(folder, chunknum[0],chunknum[1]), "a")
                    outfile.write('\n# ' + question)
                    outfile.write('\n' + answer)
                    question = ""
                    answer = ""
                    outfile.close()

                else:
                    outfile = open('%s/%s/%s.md' %(folder, chunknum[0],chunknum[1]), "w")
                    outfile.write('# ' + question)
                    outfile.write('\n' + answer)
                    question = ""
                    answer = ""
                    outfile.close()

                prev_to_prev_chunk = prev_chunk
                prev_chunk = chunknum[1]
        else:
            print "Something missing at row: %d" %r

    print "Conversion done !"
