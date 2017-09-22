from docx import Document
import os
import re

# Input file name
print "Enter file name (with file extension):"
filename = raw_input('> ')

# Initialization of variables
document = Document(filename)
table = document.tables[0]
row_count = len(table.rows)
question = ""
chunknum = ""
answer = ""
pattern = re.compile(r'-')
f = filename.split('.')
folder = f[0]

# Check if directory exists
if not os.path.exists(folder):
    os.mkdir(folder)

# Convert into md file
print "Conversion in progress !"
for r in xrange(2, row_count):

    # Read table from docx file
    row = table.rows[r]
    cell = row.cells[2]
    paragraph = cell.paragraphs[0]
    p = paragraph.text.encode('utf-8').strip()

    # Check for questions
    if p.endswith('?'):
        question = p

    # Check for answers
    elif p[-1].endswith('.'):
        answer = p

    # Check for chunk numbers
    elif p[-1].isdigit():
        if pattern.search(p):
            chunknum = p.split('-')
            if not os.path.exists('%s/%s' %(folder, chunknum[0])):
                os.mkdir("%s/%s/" %(folder, chunknum[0]))
            outfile = open('%s/%s/%s.md' %(folder, chunknum[0],chunknum[1]), "w")
            outfile.write('# ' + question)
            outfile.write('\n' + answer)
            outfile.close()
    else:
        print "Blank space found at row: %d" %r

print "Conversion done !"
